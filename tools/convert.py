from .base import *
from PySide6.QtWidgets import QDialog
from ._convert import ui
from libs.stdout import _getstamp
from libs.ui.effect import acrylic
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

tr = {
    'view' : ("已转换到 '%s'\n是否查看?", 'Converted to %s\nDo You Want to Have a View?'),
    'stamp' : ('翻译器 %s 生成于 %s', 'Generated by Translator %s on %s'),
    'error' : ('错误: %s', 'Error: %s')
}

UITr = {
    'Columns' : '列数',
    'Font' : '字体',
    'Pt' : '磅',
    'Sections' : '段落',
    'Stamp' : '时间戳',
    'Word' : '单词',
    'Word Count' : '单词计数'
}

def load_settings():
    data = tool.data.Get()
    if data:
        columns, font_family, font_size, stamp_b, stamp_format, word_count_b, word_b, info_b, trans_b = data
        ui.Columns.setValue(columns)
        ui.Font.setCurrentText(font_family)
        ui.Font_Size.setValue(font_size)
        ui.Stamp.setChecked(stamp_b)
        ui.Stamp_Format.setCurrentIndex(stamp_format)
        ui.Word_Count.setChecked(word_count_b)
        ui.Word.setChecked(word_b)
        ui.Informations.setChecked(info_b)
        ui.Translations.setChecked(trans_b)

def save_settings():
    global columns, font, font_size, stamp_b, stamp_format, word_count_b, word_b, info_b, trans_b
    columns = ui.Columns.value()
    font = ui.Font.currentText()
    font_size = ui.Font_Size.value()
    stamp_b = ui.Stamp.isChecked()
    stamp_format = ui.Stamp_Format.currentIndex()
    word_count_b = ui.Word_Count.isChecked()
    word_b = ui.Word.isChecked()
    info_b = ui.Informations.isChecked()
    trans_b = ui.Translations.isChecked()
    tool.data.Set([columns, font, font_size, stamp_b, stamp_format, word_count_b, word_b, info_b, trans_b])

def uimain():
    dialog = QDialog()
    dialog.setWindowIcon(ui.icon)
    dialog.accepted.connect(main)
    dialog.setWindowTitle(tool.get_name())
    acrylic(dialog)
    ui.setupUi(dialog, UITr)
    ui.Stamp.toggled.connect(lambda b: ui.Stamp_Format.setEnabled(b) or ui.Word_Count.setEnabled(b))
    load_settings()
    dialog.exec()

def main():
    try:
        save_settings()
        file_name = tool.dialog.SaveFile(type='*.docx')
        if file_name:
            process().save(file_name)
            if tool.message.Ask(tool.tr('view') % file_name):
                tool.dialog.Pop(file_name)
    except Exception as e:
        tool.message.Error(tool.tr('error') % e)

def process():
    from libs.config import Setting
    document = Document()
    section = document.sections[0]
    path = section._sectPr.xpath('./w:cols')[0]
    path.set(qn('w:num'), str(columns))
    path.set(qn('w:space'), '20')
    section.top_margin = Pt(10)
    section.bottom_margin = Pt(10)
    section.left_margin = Pt(20)
    section.right_margin = Pt(20)
    section.page_width = Pt(595)
    section.page_height = Pt(842)
    normal = document.styles['Normal']
    normal.font.name = font
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    normal.font.size = Pt(font_size)
    if stamp_b:
        header = section.header.paragraphs[0]
        stamp = tool.tr('stamp') % (info.version, _getstamp(stamp_format))
        header.add_run(stamp_b + f' ({tool.mw.ui.Bank.count()})' if word_count_b else '')
    for result in tool.mw.ui.Bank.results:
        information = result.phonetic
        word = ''.join(['_' if not word_b and c.isalpha() else c for c in result.word])
        head = word + f' /{information}/' if info_b and information else ''
        p = document.add_paragraph(head + f'\n{result.get_translation(Setting.Language)}' if trans_b else '')
        p.paragraph_format.line_spacing = Pt(font_size)
        p.paragraph_format.space_after = Pt(5)
    return document

tool = Tool()
tool.name = 'Convert'
tool.name_zh = '转换'
tool.doc = 'Convert vocabulary file to docx'
tool.doc_zh = '转换单词表为文档'
tool.tr.Tr = tr
tool.entrance = uimain
